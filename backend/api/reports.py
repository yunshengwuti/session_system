"""
统计相关 API
"""
from fastapi import APIRouter, Depends, Query, HTTPException, BackgroundTasks
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session as DBSession
from sqlalchemy import func
from typing import Optional
from datetime import date, timedelta, datetime
import traceback
import uuid
from io import BytesIO
from urllib.parse import quote

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from config.database import get_db, SessionLocal
from models.database import Session, Message, DailyReport, WeeklyReport, ReportTask
from models.schemas import DailyReportOut, WeeklyReportOut
from services.ai_service import generate_daily_report, generate_weekly_report

router = APIRouter(prefix="/api/reports", tags=["报告"])


DOCX_FONT = "Microsoft YaHei"
DOCX_TEXT_COLOR = RGBColor(47, 52, 55)
DOCX_MUTED_COLOR = RGBColor(96, 101, 108)


def format_datetime(value) -> str:
    if not value:
        return ""
    return value.strftime("%Y-%m-%d %H:%M:%S") if hasattr(value, "strftime") else str(value)


def apply_run_style(run, *, size: float = 10.5, bold: Optional[bool] = None, color: Optional[RGBColor] = None) -> None:
    run.font.name = DOCX_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def apply_style_font(style, *, size: float, bold: Optional[bool] = None, color: Optional[RGBColor] = None) -> None:
    style.font.name = DOCX_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color


def create_report_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    apply_style_font(normal, size=10.5, color=DOCX_TEXT_COLOR)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    apply_style_font(document.styles["Heading 1"], size=16, bold=True, color=DOCX_TEXT_COLOR)
    apply_style_font(document.styles["Heading 2"], size=13, bold=True, color=DOCX_TEXT_COLOR)
    apply_style_font(document.styles["Heading 3"], size=11.5, bold=True, color=DOCX_TEXT_COLOR)
    return document


def add_report_heading(document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(str(text), level=level)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(14)
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(4)

    sizes = {1: 16, 2: 13, 3: 11.5}
    for run in paragraph.runs:
        apply_run_style(run, size=sizes.get(level, 10.5), bold=True, color=DOCX_TEXT_COLOR)


def add_body_paragraph(document, text: object) -> None:
    paragraph = document.add_paragraph("" if text is None else str(text))
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        apply_run_style(run, size=10.5, color=DOCX_TEXT_COLOR)


def format_table_cell(cell, *, is_header: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = 1.2
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            apply_run_style(
                run,
                size=10.5 if is_header else 10,
                bold=is_header,
                color=DOCX_MUTED_COLOR if is_header else DOCX_TEXT_COLOR,
            )


def add_table(document, headers: list[str], rows: list[list]) -> None:
    if not rows:
        add_body_paragraph(document, "暂无数据")
        return

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
        format_table_cell(table.rows[0].cells[index], is_header=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = "" if value is None else str(value)
            format_table_cell(cells[index])

    document.add_paragraph()


def add_key_value_table(document, rows: list[tuple[str, object]]) -> None:
    add_table(document, ["项目", "内容"], rows)


def add_dict_table(document, title: str, data: Optional[dict], value_label: str = "次数") -> None:
    add_report_heading(document, title, level=2)
    rows = [[key, value] for key, value in (data or {}).items()]
    add_table(document, ["名称", value_label], rows)


def build_docx_response(document, filename: str) -> StreamingResponse:
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    encoded_filename = quote(filename)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


def task_payload(task: ReportTask) -> dict:
    return {
        "task_id": task.task_id,
        "report_type": task.report_type,
        "target_key": task.target_key,
        "status": task.status,
        "progress": task.progress or 0,
        "message": task.message or "",
        "error": task.error,
        "result_key": task.result_key,
        "created_at": task.created_at,
        "updated_at": task.updated_at,
    }


def update_task(
    db: DBSession,
    task_id: str,
    *,
    status: Optional[str] = None,
    progress: Optional[int] = None,
    message: Optional[str] = None,
    error: Optional[str] = None,
    result_key: Optional[str] = None,
) -> None:
    task = db.get(ReportTask, task_id)
    if not task:
        return
    if status is not None:
        task.status = status
    if progress is not None:
        task.progress = max(0, min(100, int(progress)))
    if message is not None:
        task.message = message[:500]
    if error is not None:
        task.error = error
    if result_key is not None:
        task.result_key = result_key
    task.updated_at = datetime.now()
    db.commit()


def build_daily_sessions_data(db: DBSession, report_date: date) -> list[dict]:
    sessions = db.query(Session).filter(Session.session_date == report_date).all()
    sessions_data = []
    for s in sessions:
        messages = db.query(Message).filter(Message.session_id == s.session_id).order_by(Message.message_time).all()
        messages_list = [
            {
                "speaker": m.speaker,
                "content": m.content,
                "message_type": m.message_type,
                "image_url": m.image_url,
            }
            for m in messages
        ]

        sessions_data.append({
            "session_id": s.session_id,
            "customer_name": s.customer_name,
            "org_name": s.org_name,
            "customer_service": s.customer_service,
            "duration_seconds": s.duration_seconds,
            "session_date": s.session_date,
            "messages": messages_list,
        })
    return sessions_data


def get_workdays(start_date: date, end_date: date) -> list[date]:
    delta_days = (end_date - start_date).days + 1
    if not (3 <= delta_days <= 7):
        raise ValueError("时间段必须为3-7天")

    workdays = []
    for i in range(delta_days):
        day = start_date + timedelta(days=i)
        if day.weekday() < 5:
            workdays.append(day)

    if len(workdays) < 3:
        raise ValueError(f"时间段内工作日不足3天（当前{len(workdays)}天），请重新选择")
    if len(workdays) > 5:
        raise ValueError(f"时间段内工作日超过5天（当前{len(workdays)}天），建议选择周一到周五")
    return workdays


async def run_daily_report_task(task_id: str, report_date_str: str) -> None:
    db = SessionLocal()
    try:
        report_date = date.fromisoformat(report_date_str)
        update_task(db, task_id, status="running", progress=3, message="正在检查日报数据")

        existing = db.query(DailyReport).filter(DailyReport.report_date == report_date).first()
        if existing:
            update_task(db, task_id, status="succeeded", progress=100, message="日报已存在", result_key=str(report_date))
            return

        sessions_data = build_daily_sessions_data(db, report_date)
        if not sessions_data:
            raise ValueError(f"{report_date} 没有会话数据")

        async def progress_callback(progress: int, message: str) -> None:
            update_task(db, task_id, status="running", progress=progress, message=message)

        report_result = await generate_daily_report(sessions_data, progress_callback=progress_callback)

        daily_report = DailyReport(
            report_date=report_date,
            total_sessions=len(sessions_data),
            keywords_json=report_result.get("keywords_json"),
            category_stats_json=report_result.get("category_stats_json"),
            long_duration_issues=report_result.get("long_duration_issues"),
            org_distribution_json=report_result.get("org_distribution_json"),
            service_stats_json=report_result.get("service_stats_json"),
            ai_summary=report_result.get("ai_summary"),
            generated_at=datetime.now(),
        )
        db.add(daily_report)
        db.commit()
        update_task(db, task_id, status="succeeded", progress=100, message="日报生成完成", result_key=str(report_date))
    except Exception as exc:
        traceback.print_exc()
        update_task(db, task_id, status="failed", progress=100, message="日报生成失败", error=str(exc))
    finally:
        db.close()


async def run_weekly_report_task(task_id: str, start_date_str: str, end_date_str: str) -> None:
    db = SessionLocal()
    try:
        start_date = date.fromisoformat(start_date_str)
        end_date = date.fromisoformat(end_date_str)
        update_task(db, task_id, status="running", progress=3, message="正在检查周报数据")

        existing = db.query(WeeklyReport).filter(
            WeeklyReport.week_start_date == start_date,
            WeeklyReport.week_end_date == end_date
        ).first()
        if existing:
            update_task(db, task_id, status="succeeded", progress=100, message="周报已存在", result_key=str(start_date))
            return

        workdays = get_workdays(start_date, end_date)
        existing_reports = {}
        missing_dates = []
        for day in workdays:
            daily_report = db.query(DailyReport).filter(DailyReport.report_date == day).first()
            if daily_report:
                existing_reports[day] = daily_report
            else:
                missing_dates.append(day)

        update_task(db, task_id, status="running", progress=10, message="正在补齐缺失日报")
        for index, missing_date in enumerate(missing_dates):
            sessions_data = build_daily_sessions_data(db, missing_date)
            if not sessions_data:
                raise ValueError(f"工作日 {missing_date} 没有会话数据，无法生成周报")

            async def daily_progress_callback(progress: int, message: str, index=index, missing_date=missing_date) -> None:
                overall = 10 + int(60 * ((index + progress / 100) / max(len(missing_dates), 1)))
                update_task(db, task_id, status="running", progress=overall, message=f"{missing_date}：{message}")

            report_result = await generate_daily_report(sessions_data, progress_callback=daily_progress_callback)
            daily_report = DailyReport(
                report_date=missing_date,
                total_sessions=len(sessions_data),
                keywords_json=report_result.get("keywords_json"),
                category_stats_json=report_result.get("category_stats_json"),
                long_duration_issues=report_result.get("long_duration_issues"),
                org_distribution_json=report_result.get("org_distribution_json"),
                service_stats_json=report_result.get("service_stats_json"),
                ai_summary=report_result.get("ai_summary", ""),
                generated_at=datetime.now(),
            )
            db.add(daily_report)
            db.commit()
            db.refresh(daily_report)
            existing_reports[missing_date] = daily_report

        daily_reports_data = []
        for day in workdays:
            daily_report = existing_reports[day]
            daily_reports_data.append({
                "report_date": daily_report.report_date,
                "total_sessions": daily_report.total_sessions,
                "keywords_json": daily_report.keywords_json,
                "category_stats_json": daily_report.category_stats_json,
                "long_duration_issues": daily_report.long_duration_issues,
                "org_distribution_json": daily_report.org_distribution_json,
                "service_stats_json": daily_report.service_stats_json,
            })

        async def weekly_progress_callback(progress: int, message: str) -> None:
            update_task(db, task_id, status="running", progress=progress, message=message)

        update_task(db, task_id, status="running", progress=78, message="正在生成周报")
        weekly_result = await generate_weekly_report([], daily_reports_data, progress_callback=weekly_progress_callback)
        weekly_report = WeeklyReport(
            week_start_date=start_date,
            week_end_date=end_date,
            total_sessions=sum([r["total_sessions"] for r in daily_reports_data]),
            keywords_json=weekly_result.get("keywords_json"),
            category_stats_json=weekly_result.get("category_stats_json"),
            org_distribution_json=weekly_result.get("org_distribution_json"),
            service_stats_json=weekly_result.get("service_stats_json"),
            daily_trend_json=weekly_result.get("daily_trend_json"),
            ai_summary=weekly_result.get("ai_summary"),
            generated_at=datetime.now(),
        )
        db.add(weekly_report)
        db.commit()
        update_task(db, task_id, status="succeeded", progress=100, message="周报生成完成", result_key=str(start_date))
    except Exception as exc:
        traceback.print_exc()
        update_task(db, task_id, status="failed", progress=100, message="周报生成失败", error=str(exc))
    finally:
        db.close()




@router.get("/tasks/{task_id}")
def get_report_task(task_id: str, db: DBSession = Depends(get_db)):
    task = db.get(ReportTask, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return task_payload(task)

# 列表路由必须在参数路由之前
@router.get("/daily/list")
def list_daily_reports(
    date_from: Optional[date] = Query(None, description="开始日期"),
    date_to: Optional[date] = Query(None, description="结束日期"),
    db: DBSession = Depends(get_db),
):
    """获取日报列表"""
    query = db.query(DailyReport)

    if date_from:
        query = query.filter(DailyReport.report_date >= date_from)
    if date_to:
        query = query.filter(DailyReport.report_date <= date_to)

    reports = query.order_by(DailyReport.report_date.desc()).all()
    return {"reports": reports}


@router.get("/weekly/list")
def list_weekly_reports(
    db: DBSession = Depends(get_db),
):
    """获取周报列表"""
    reports = db.query(WeeklyReport).order_by(WeeklyReport.week_start_date.desc()).all()
    return reports


@router.post("/daily/task")
def create_daily_report_task(
    background_tasks: BackgroundTasks,
    report_date: date = Query(..., description="报告日期"),
    db: DBSession = Depends(get_db),
):
    """启动日报后台生成任务"""
    existing_report = db.query(DailyReport).filter(DailyReport.report_date == report_date).first()
    if existing_report:
        raise HTTPException(status_code=400, detail=f"{report_date} 的日报已存在，如需重新生成请先删除")

    existing_task = db.query(ReportTask).filter(
        ReportTask.report_type == "daily",
        ReportTask.target_key == str(report_date),
        ReportTask.status.in_(["pending", "running"])
    ).order_by(ReportTask.created_at.desc()).first()
    if existing_task:
        return task_payload(existing_task)

    task = ReportTask(
        task_id=uuid.uuid4().hex,
        report_type="daily",
        target_key=str(report_date),
        status="pending",
        progress=0,
        message="日报任务已创建",
    )
    db.add(task)
    db.commit()
    db.refresh(task)
    background_tasks.add_task(run_daily_report_task, task.task_id, str(report_date))
    return task_payload(task)


@router.post("/weekly/task")
def create_weekly_report_task(
    background_tasks: BackgroundTasks,
    start_date: date = Query(..., description="周报开始日期"),
    end_date: date = Query(..., description="周报结束日期"),
    db: DBSession = Depends(get_db),
):
    """启动周报后台生成任务"""
    try:
        get_workdays(start_date, end_date)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    existing_report = db.query(WeeklyReport).filter(
        WeeklyReport.week_start_date == start_date,
        WeeklyReport.week_end_date == end_date
    ).first()
    if existing_report:
        raise HTTPException(status_code=400, detail=f"{start_date} ~ {end_date} 的周报已存在，如需重新生成请先删除")

    target_key = f"{start_date}:{end_date}"
    existing_task = db.query(ReportTask).filter(
        ReportTask.report_type == "weekly",
        ReportTask.target_key == target_key,
        ReportTask.status.in_(["pending", "running"])
    ).order_by(ReportTask.created_at.desc()).first()
    if existing_task:
        return task_payload(existing_task)

    task = ReportTask(
        task_id=uuid.uuid4().hex,
        report_type="weekly",
        target_key=target_key,
        status="pending",
        progress=0,
        message="周报任务已创建",
    )
    db.add(task)
    db.commit()
    db.refresh(task)
    background_tasks.add_task(run_weekly_report_task, task.task_id, str(start_date), str(end_date))
    return task_payload(task)


@router.post("/daily", response_model=DailyReportOut)
async def create_daily_report(
    report_date: date = Query(..., description="报告日期"),
    db: DBSession = Depends(get_db),
):
    """生成某天的AI日报"""
    # 检查是否已存在
    existing = db.query(DailyReport).filter(DailyReport.report_date == report_date).first()
    if existing:
        raise HTTPException(status_code=400, detail=f"{report_date} 的日报已存在，如需重新生成请先删除")

    # 获取当天所有会话数据
    sessions = db.query(Session).filter(Session.session_date == report_date).all()

    if not sessions:
        raise HTTPException(status_code=404, detail=f"{report_date} 没有会话数据")

    # 准备数据给AI
    sessions_data = []
    for s in sessions:
        messages = db.query(Message).filter(Message.session_id == s.session_id).order_by(Message.message_time).all()
        messages_list = [
            {
                "speaker": m.speaker,
                "content": m.content,
                "message_type": m.message_type,
                "image_url": m.image_url,
            }
            for m in messages
        ]

        sessions_data.append({
            "session_id": s.session_id,
            "customer_name": s.customer_name,
            "org_name": s.org_name,
            "customer_service": s.customer_service,
            "duration_seconds": s.duration_seconds,
            "messages": messages_list,
        })

    # 调用AI生成报告
    report_result = await generate_daily_report(sessions_data)

    # 保存到数据库
    daily_report = DailyReport(
        report_date=report_date,
        total_sessions=len(sessions),
        keywords_json=report_result.get("keywords_json"),
        category_stats_json=report_result.get("category_stats_json"),
        long_duration_issues=report_result.get("long_duration_issues"),
        org_distribution_json=report_result.get("org_distribution_json"),
        service_stats_json=report_result.get("service_stats_json"),
        ai_summary=report_result.get("ai_summary"),
        generated_at=datetime.now(),
    )
    db.add(daily_report)
    db.commit()
    db.refresh(daily_report)

    return daily_report



@router.get("/daily/{report_date}/export")
def export_daily_report(
    report_date: date,
    db: DBSession = Depends(get_db),
):
    """导出某天日报为 Word"""
    report = db.query(DailyReport).filter(DailyReport.report_date == report_date).first()
    if not report:
        raise HTTPException(status_code=404, detail=f"{report_date} 的日报不存在")

    document = create_report_document()
    add_report_heading(document, f"日报 - {report.report_date}", level=1)

    add_report_heading(document, "基本信息", level=2)
    add_key_value_table(document, [
        ("日期", report.report_date),
        ("总会话数", report.total_sessions),
        ("生成时间", format_datetime(report.generated_at)),
    ])

    add_report_heading(document, "今日问题概览", level=2)
    add_body_paragraph(document, report.long_duration_issues or "暂无内容")

    add_dict_table(document, "关键问题统计", report.keywords_json, "次数")
    add_dict_table(document, "问题业务分类", report.category_stats_json, "占比/数值")

    org_data = report.org_distribution_json or {}
    risks = org_data.get("risks", [])
    add_report_heading(document, "异常与风险", level=2)
    add_table(
        document,
        ["风险类型", "问题描述", "影响客户数", "紧急程度", "建议处理方式"],
        [[
            risk.get("risk_type", ""),
            risk.get("description", ""),
            risk.get("affected_customers", ""),
            risk.get("urgency", ""),
            risk.get("suggestion", ""),
        ] for risk in risks]
    )

    suggestions = org_data.get("suggestions", [])
    add_report_heading(document, "行动建议", level=2)
    if suggestions:
        for index, suggestion in enumerate(suggestions, start=1):
            add_body_paragraph(document, f"{index}. {suggestion}")
    else:
        add_body_paragraph(document, "暂无建议")

    add_dict_table(document, "客服工作量", report.service_stats_json, "会话数")

    return build_docx_response(document, f"日报_{report.report_date}.docx")

@router.get("/daily/{report_date}", response_model=DailyReportOut)
def get_daily_report(
    report_date: date,
    db: DBSession = Depends(get_db),
):
    """获取某天的日报"""
    report = db.query(DailyReport).filter(DailyReport.report_date == report_date).first()
    if not report:
        raise HTTPException(status_code=404, detail=f"{report_date} 的日报不存在")
    return report


@router.delete("/daily/{report_date}")
def delete_daily_report(
    report_date: date,
    db: DBSession = Depends(get_db),
):
    """删除某天的日报"""
    report = db.query(DailyReport).filter(DailyReport.report_date == report_date).first()
    if not report:
        raise HTTPException(status_code=404, detail=f"{report_date} 的日报不存在")

    db.delete(report)
    db.commit()
    return {"message": f"{report_date} 的日报已删除"}


@router.post("/weekly", response_model=WeeklyReportOut)
async def create_weekly_report(
    start_date: date = Query(..., description="周报开始日期"),
    end_date: date = Query(..., description="周报结束日期"),
    db: DBSession = Depends(get_db),
):
    """生成周报（自动剔除周末，工作日3-5天）"""

    # 1. 验证时间段
    delta_days = (end_date - start_date).days + 1
    if not (3 <= delta_days <= 7):
        raise HTTPException(status_code=400, detail="时间段必须为3-7天")

    # 2. 获取时间段内的所有工作日（周一到周五）
    workdays = []
    for i in range(delta_days):
        day = start_date + timedelta(days=i)
        # 0=周一, 1=周二, ..., 4=周五, 5=周六, 6=周日
        if day.weekday() < 5:  # 工作日
            workdays.append(day)

    # 验证工作日数量
    if len(workdays) < 3:
        raise HTTPException(
            status_code=400,
            detail=f"时间段内工作日不足3天（当前{len(workdays)}天），请重新选择"
        )

    if len(workdays) > 5:
        raise HTTPException(
            status_code=400,
            detail=f"时间段内工作日超过5天（当前{len(workdays)}天），建议选择周一到周五"
        )

    print(f"\n📅 选择的时间段：{start_date} ~ {end_date}")
    print(f"📅 工作日：{workdays} (共{len(workdays)}天)\n")

    # 检查是否已存在
    existing = db.query(WeeklyReport).filter(
        WeeklyReport.week_start_date == start_date,
        WeeklyReport.week_end_date == end_date
    ).first()
    if existing:
        raise HTTPException(
            status_code=400,
            detail=f"{start_date} ~ {end_date} 的周报已存在，如需重新生成请先删除"
        )

    # 3. 检查并补齐工作日的日报
    missing_dates = []
    existing_reports = {}

    for day in workdays:
        daily_report = db.query(DailyReport).filter(DailyReport.report_date == day).first()

        if daily_report:
            existing_reports[day] = daily_report
            print(f"✅ {day} 日报已存在")
        else:
            missing_dates.append(day)

    # 4. 生成缺失的日报
    if missing_dates:
        print(f"\n📋 需要生成 {len(missing_dates)} 份日报：{missing_dates}\n")

        for missing_date in missing_dates:
            # 获取该日期的会话数据
            sessions = db.query(Session).filter(
                Session.session_date == missing_date
            ).all()

            if not sessions:
                raise HTTPException(
                    status_code=404,
                    detail=f"工作日 {missing_date} 没有会话数据，无法生成周报"
                )

            # 准备会话数据
            sessions_data = []
            for s in sessions:
                messages = db.query(Message).filter(
                    Message.session_id == s.session_id
                ).order_by(Message.message_time).all()

                messages_list = [
                    {
                        "speaker": m.speaker,
                        "content": m.content,
                        "message_type": m.message_type,
                    }
                    for m in messages
                ]

                sessions_data.append({
                    "session_id": s.session_id,
                    "customer_name": s.customer_name,
                    "org_name": s.org_name,
                    "customer_service": s.customer_service,
                    "duration_seconds": s.duration_seconds,
                    "session_date": s.session_date,
                    "messages": messages_list,
                })

            # 生成日报
            print(f"📊 生成 {missing_date} 的日报（共{len(sessions_data)}个会话）...")
            report_result = await generate_daily_report(sessions_data)

            # 保存日报
            daily_report = DailyReport(
                report_date=missing_date,
                total_sessions=len(sessions_data),
                keywords_json=report_result["keywords_json"],
                category_stats_json=report_result["category_stats_json"],
                long_duration_issues=report_result["long_duration_issues"],
                org_distribution_json=report_result["org_distribution_json"],
                service_stats_json=report_result["service_stats_json"],
                ai_summary=report_result.get("ai_summary", "")
            )
            db.add(daily_report)
            db.commit()
            db.refresh(daily_report)

            existing_reports[missing_date] = daily_report
            print(f"✅ {missing_date} 日报生成完成\n")

    # 5. 收集所有工作日的日报数据
    daily_reports_data = []
    for day in workdays:
        daily_report = existing_reports[day]
        daily_reports_data.append({
            "report_date": daily_report.report_date,
            "total_sessions": daily_report.total_sessions,
            "keywords_json": daily_report.keywords_json,
            "category_stats_json": daily_report.category_stats_json,
            "long_duration_issues": daily_report.long_duration_issues,
            "org_distribution_json": daily_report.org_distribution_json,
            "service_stats_json": daily_report.service_stats_json,
        })

    # 6. 生成周报
    print(f"\n📊 开始生成周报（{start_date} ~ {end_date}，{len(workdays)}个工作日）...\n")
    weekly_result = await generate_weekly_report([], daily_reports_data)

    # 7. 保存周报
    weekly_report = WeeklyReport(
        week_start_date=start_date,
        week_end_date=end_date,
        total_sessions=sum([r["total_sessions"] for r in daily_reports_data]),
        keywords_json=weekly_result["keywords_json"],
        category_stats_json=weekly_result["category_stats_json"],
        org_distribution_json=weekly_result["org_distribution_json"],
        service_stats_json=weekly_result["service_stats_json"],
        daily_trend_json=weekly_result["daily_trend_json"],
        ai_summary=weekly_result["ai_summary"]
    )
    db.add(weekly_report)
    db.commit()
    db.refresh(weekly_report)

    print(f"✅ 周报生成完成\n")

    return weekly_report

    # 调用AI生成周报
    report_result = await generate_weekly_report(week_sessions_data, daily_reports_data)

    # 保存到数据库
    weekly_report = WeeklyReport(
        week_start_date=week_start_date,
        week_end_date=week_end_date,
        total_sessions=len(sessions),
        keywords_json=report_result.get("keywords_json"),
        category_stats_json=report_result.get("category_stats_json"),
        org_distribution_json=report_result.get("org_distribution_json"),
        service_stats_json=report_result.get("service_stats_json"),
        daily_trend_json=report_result.get("daily_trend_json"),
        ai_summary=report_result.get("ai_summary"),
        generated_at=datetime.now(),
    )
    db.add(weekly_report)
    db.commit()
    db.refresh(weekly_report)

    return weekly_report



@router.get("/weekly/{week_start_date}/export")
def export_weekly_report(
    week_start_date: date,
    db: DBSession = Depends(get_db),
):
    """导出某周周报为 Word"""
    report = db.query(WeeklyReport).filter(WeeklyReport.week_start_date == week_start_date).first()
    if not report:
        raise HTTPException(status_code=404, detail=f"{week_start_date} 开始的周报不存在")

    document = create_report_document()
    add_report_heading(document, f"周报 - {report.week_start_date} 至 {report.week_end_date}", level=1)

    add_report_heading(document, "基本信息", level=2)
    add_key_value_table(document, [
        ("周期", f"{report.week_start_date} 至 {report.week_end_date}"),
        ("总会话数", report.total_sessions),
        ("生成时间", format_datetime(report.generated_at)),
    ])

    add_report_heading(document, "本周总结", level=2)
    add_body_paragraph(document, report.ai_summary or "暂无内容")

    add_dict_table(document, "高频问题TOP", report.keywords_json, "次数")
    add_dict_table(document, "问题业务分类", report.category_stats_json, "占比/数值")

    add_report_heading(document, "每日趋势", level=2)
    add_table(
        document,
        ["日期", "会话数", "摘要"],
        [[item.get("date", ""), item.get("sessions", ""), item.get("brief", "")] for item in (report.daily_trend_json or [])]
    )

    org_data = report.org_distribution_json or {}
    add_report_heading(document, "本周趋势", level=2)
    add_body_paragraph(document, org_data.get("trends", "暂无内容"))

    add_report_heading(document, "重点风险", level=2)
    add_table(
        document,
        ["风险类型", "问题描述", "建议"],
        [[
            risk.get("risk_type", ""),
            risk.get("description", ""),
            risk.get("suggestion", ""),
        ] for risk in org_data.get("key_risks", [])]
    )

    cases = org_data.get("cases", [])
    add_report_heading(document, "典型案例", level=2)
    if cases:
        for index, case_item in enumerate(cases, start=1):
            add_report_heading(document, f"案例{index}：{case_item.get('title', '')}", level=3)
            add_body_paragraph(document, case_item.get("description", ""))
            outcome = case_item.get("outcome")
            if outcome:
                add_body_paragraph(document, f"处理结果：{outcome}")
    else:
        add_body_paragraph(document, "暂无案例")

    next_week_plan = org_data.get("next_week_plan", [])
    add_report_heading(document, "下周改进计划", level=2)
    if next_week_plan:
        for index, plan in enumerate(next_week_plan, start=1):
            add_body_paragraph(document, f"{index}. {plan}")
    else:
        add_body_paragraph(document, "暂无计划")

    add_dict_table(document, "客服工作量", report.service_stats_json, "会话数")

    return build_docx_response(document, f"周报_{report.week_start_date}_{report.week_end_date}.docx")

@router.get("/weekly/{week_start_date}", response_model=WeeklyReportOut)
def get_weekly_report(
    week_start_date: date,
    db: DBSession = Depends(get_db),
):
    """获取某周的周报"""
    report = db.query(WeeklyReport).filter(WeeklyReport.week_start_date == week_start_date).first()
    if not report:
        raise HTTPException(status_code=404, detail=f"{week_start_date} 开始的周报不存在")
    return report


@router.delete("/weekly/{week_start_date}")
def delete_weekly_report(
    week_start_date: date,
    db: DBSession = Depends(get_db),
):
    """删除某周的周报"""
    report = db.query(WeeklyReport).filter(WeeklyReport.week_start_date == week_start_date).first()
    if not report:
        raise HTTPException(status_code=404, detail=f"{week_start_date} 开始的周报不存在")

    db.delete(report)
    db.commit()
    return {"message": f"{week_start_date} 开始的周报已删除"}
